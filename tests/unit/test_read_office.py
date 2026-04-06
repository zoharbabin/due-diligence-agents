"""Tests for the read_office MCP tool."""

from __future__ import annotations

from typing import TYPE_CHECKING

import pytest

if TYPE_CHECKING:
    from pathlib import Path


class TestReadExcel:
    """Tests for reading Excel files (.xlsx)."""

    def test_read_xlsx_single_sheet(self, tmp_path: Path) -> None:
        """Read a simple .xlsx with one sheet returns markdown table."""
        openpyxl = pytest.importorskip("openpyxl")
        wb = openpyxl.Workbook()
        ws = wb.active
        assert ws is not None
        ws.title = "Data"
        ws.append(["Name", "Value"])
        ws.append(["Alice", 100])
        ws.append(["Bob", 200])
        wb.save(tmp_path / "test.xlsx")

        from dd_agents.tools.read_office import read_office

        result = read_office(str(tmp_path / "test.xlsx"))
        assert result["status"] == "ok"
        content = result["content"]
        assert "Sheet: Data" in content
        assert "Alice" in content
        assert "Bob" in content
        assert "100" in content

    def test_read_xlsx_multiple_sheets(self, tmp_path: Path) -> None:
        """Read a multi-sheet .xlsx returns all sheets."""
        openpyxl = pytest.importorskip("openpyxl")
        wb = openpyxl.Workbook()
        ws1 = wb.active
        assert ws1 is not None
        ws1.title = "Revenue"
        ws1.append(["Q1", "Q2"])
        ws1.append([1000, 2000])

        ws2 = wb.create_sheet("Costs")
        ws2.append(["Item", "Amount"])
        ws2.append(["Rent", 500])
        wb.save(tmp_path / "multi.xlsx")

        from dd_agents.tools.read_office import read_office

        result = read_office(str(tmp_path / "multi.xlsx"))
        assert result["status"] == "ok"
        assert "Sheet: Revenue" in result["content"]
        assert "Sheet: Costs" in result["content"]
        assert "1000" in result["content"]
        assert "Rent" in result["content"]

    def test_read_xlsx_specific_sheet(self, tmp_path: Path) -> None:
        """Reading with sheet_name filters to that sheet only."""
        openpyxl = pytest.importorskip("openpyxl")
        wb = openpyxl.Workbook()
        ws1 = wb.active
        assert ws1 is not None
        ws1.title = "Summary"
        ws1.append(["Total", 999])

        ws2 = wb.create_sheet("Detail")
        ws2.append(["Item", "Cost"])
        ws2.append(["Widget", 42])
        wb.save(tmp_path / "sheets.xlsx")

        from dd_agents.tools.read_office import read_office

        result = read_office(str(tmp_path / "sheets.xlsx"), sheet_name="Detail")
        assert result["status"] == "ok"
        assert "Widget" in result["content"]
        assert "999" not in result["content"]

    def test_read_xlsx_invalid_sheet_name(self, tmp_path: Path) -> None:
        """Invalid sheet_name returns error immediately (no fallback attempted)."""
        openpyxl = pytest.importorskip("openpyxl")
        wb = openpyxl.Workbook()
        ws = wb.active
        assert ws is not None
        ws.title = "Data"
        ws.append(["A", 1])
        wb.save(tmp_path / "test.xlsx")

        from dd_agents.tools.read_office import read_office

        result = read_office(str(tmp_path / "test.xlsx"), sheet_name="NonExistent")
        assert result["status"] == "error"
        assert "not found" in result["reason"].lower()

    def test_read_xlsx_empty_sheet(self, tmp_path: Path) -> None:
        """Empty Excel file returns informative message."""
        openpyxl = pytest.importorskip("openpyxl")
        wb = openpyxl.Workbook()
        wb.save(tmp_path / "empty.xlsx")

        from dd_agents.tools.read_office import read_office

        result = read_office(str(tmp_path / "empty.xlsx"))
        assert result["status"] == "ok"
        assert "Sheet" in result["content"] or "empty" in result["content"].lower()

    def test_read_xlsx_preserves_dates(self, tmp_path: Path) -> None:
        """String dates are preserved as-is."""
        openpyxl = pytest.importorskip("openpyxl")
        wb = openpyxl.Workbook()
        ws = wb.active
        assert ws is not None
        ws.append(["Date", "Event"])
        ws.append(["2024-01-15", "Launch"])
        wb.save(tmp_path / "dates.xlsx")

        from dd_agents.tools.read_office import read_office

        result = read_office(str(tmp_path / "dates.xlsx"))
        assert result["status"] == "ok"
        assert "2024-01-15" in result["content"]

    def test_read_xlsx_uses_column_letter_headers(self, tmp_path: Path) -> None:
        """Table headers are Excel column letters (A, B, C), not data values."""
        openpyxl = pytest.importorskip("openpyxl")
        wb = openpyxl.Workbook()
        ws = wb.active
        assert ws is not None
        ws.append(["Revenue", "$1M"])
        ws.append(["Costs", "$500K"])
        wb.save(tmp_path / "data.xlsx")

        from dd_agents.tools.read_office import read_office

        result = read_office(str(tmp_path / "data.xlsx"))
        assert result["status"] == "ok"
        content = result["content"]
        # Column letters as headers, not first row values
        assert "| A " in content
        assert "| B " in content
        # All data rows present (not skipped as header)
        assert "Revenue" in content
        assert "$1M" in content
        assert "Costs" in content

    def test_read_xlsx_pipe_in_cell_escaped(self, tmp_path: Path) -> None:
        """Pipe characters in cell values are escaped for markdown tables."""
        openpyxl = pytest.importorskip("openpyxl")
        wb = openpyxl.Workbook()
        ws = wb.active
        assert ws is not None
        ws.append(["Option A | Option B", "Yes"])
        wb.save(tmp_path / "pipes.xlsx")

        from dd_agents.tools.read_office import read_office

        result = read_office(str(tmp_path / "pipes.xlsx"))
        assert result["status"] == "ok"
        assert "\\|" in result["content"]
        # The table should still be valid — 2 data columns, not 3
        lines = [ln for ln in result["content"].split("\n") if ln.startswith("|")]
        # Header row + separator + 1 data row = 3 lines
        assert len(lines) == 3

    def test_read_xlsx_newline_in_cell_sanitized(self, tmp_path: Path) -> None:
        """Newlines in cell values are replaced with spaces for valid markdown."""
        openpyxl = pytest.importorskip("openpyxl")
        wb = openpyxl.Workbook()
        ws = wb.active
        assert ws is not None
        ws.append(["Line1\nLine2", "OK"])
        ws.append(["Normal", "Also\r\nfine"])
        wb.save(tmp_path / "newlines.xlsx")

        from dd_agents.tools.read_office import read_office

        result = read_office(str(tmp_path / "newlines.xlsx"))
        assert result["status"] == "ok"
        # Newlines must not appear in table rows (they'd break markdown)
        table_lines = [ln for ln in result["content"].split("\n") if ln.startswith("|")]
        # Header + separator + 2 data rows = 4 lines
        assert len(table_lines) == 4
        # Content is preserved (newlines become spaces)
        assert "Line1 Line2" in result["content"]
        assert "Also fine" in result["content"]


class TestCellFormatting:
    """Tests for E-1 (date→ISO-8601), E-2 (currency/percentage), E-3 (sub-tables)."""

    # --- E-1: datetime → ISO-8601 ---

    def test_datetime_cells_render_as_iso_date(self, tmp_path: Path) -> None:
        """datetime cells with no time component render as YYYY-MM-DD."""
        openpyxl = pytest.importorskip("openpyxl")
        from datetime import datetime

        wb = openpyxl.Workbook()
        ws = wb.active
        assert ws is not None
        ws.cell(row=1, column=1, value="Date")
        ws.cell(row=1, column=2, value="Event")
        ws.cell(row=2, column=1, value=datetime(2024, 3, 15))
        ws.cell(row=2, column=2, value="Launch")
        wb.save(tmp_path / "dates.xlsx")

        from dd_agents.tools.read_office import read_office

        result = read_office(str(tmp_path / "dates.xlsx"))
        assert result["status"] == "ok"
        content = result["content"]
        # Must be clean ISO, not "2024-03-15 00:00:00"
        assert "2024-03-15" in content
        assert "00:00:00" not in content

    def test_datetime_with_time_preserves_hours_minutes(self, tmp_path: Path) -> None:
        """datetime cells with non-zero time render as YYYY-MM-DD HH:MM."""
        openpyxl = pytest.importorskip("openpyxl")
        from datetime import datetime

        wb = openpyxl.Workbook()
        ws = wb.active
        assert ws is not None
        ws.cell(row=1, column=1, value=datetime(2024, 6, 1, 14, 30))
        wb.save(tmp_path / "time.xlsx")

        from dd_agents.tools.read_office import read_office

        result = read_office(str(tmp_path / "time.xlsx"))
        assert result["status"] == "ok"
        assert "2024-06-01 14:30" in result["content"]

    def test_date_object_renders_as_iso(self, tmp_path: Path) -> None:
        """Pure date objects (no time) render as YYYY-MM-DD."""
        openpyxl = pytest.importorskip("openpyxl")
        from datetime import date

        wb = openpyxl.Workbook()
        ws = wb.active
        assert ws is not None
        ws.cell(row=1, column=1, value=date(2025, 12, 25))
        wb.save(tmp_path / "date_only.xlsx")

        from dd_agents.tools.read_office import read_office

        result = read_office(str(tmp_path / "date_only.xlsx"))
        assert result["status"] == "ok"
        assert "2025-12-25" in result["content"]

    # --- E-2: Currency formatting ---

    def test_currency_formatted_cells(self, tmp_path: Path) -> None:
        """Cells with currency number_format render with $ and commas."""
        openpyxl = pytest.importorskip("openpyxl")
        wb = openpyxl.Workbook()
        ws = wb.active
        assert ws is not None
        ws.cell(row=1, column=1, value="Item")
        ws.cell(row=1, column=2, value="Price")
        cell = ws.cell(row=2, column=1, value="Widget")
        cell = ws.cell(row=2, column=2, value=1234567.89)
        cell.number_format = "$#,##0.00"
        wb.save(tmp_path / "currency.xlsx")

        from dd_agents.tools.read_office import read_office

        result = read_office(str(tmp_path / "currency.xlsx"))
        assert result["status"] == "ok"
        assert "$1,234,567.89" in result["content"]

    def test_euro_formatted_cells(self, tmp_path: Path) -> None:
        """Euro-formatted cells use the € symbol."""
        openpyxl = pytest.importorskip("openpyxl")
        wb = openpyxl.Workbook()
        ws = wb.active
        assert ws is not None
        cell = ws.cell(row=1, column=1, value=5000.00)
        cell.number_format = "€#,##0.00"
        wb.save(tmp_path / "euro.xlsx")

        from dd_agents.tools.read_office import read_office

        result = read_office(str(tmp_path / "euro.xlsx"))
        assert result["status"] == "ok"
        assert "€5,000.00" in result["content"]

    # --- E-2: Percentage formatting ---

    def test_percentage_formatted_cells(self, tmp_path: Path) -> None:
        """Cells with percentage format render as X.X% (value × 100)."""
        openpyxl = pytest.importorskip("openpyxl")
        wb = openpyxl.Workbook()
        ws = wb.active
        assert ws is not None
        ws.cell(row=1, column=1, value="Metric")
        ws.cell(row=1, column=2, value="Value")
        cell = ws.cell(row=2, column=1, value="Margin")
        cell = ws.cell(row=2, column=2, value=0.452)
        cell.number_format = "0.0%"
        wb.save(tmp_path / "pct.xlsx")

        from dd_agents.tools.read_office import read_office

        result = read_office(str(tmp_path / "pct.xlsx"))
        assert result["status"] == "ok"
        assert "45.2%" in result["content"]

    def test_whole_number_percentage(self, tmp_path: Path) -> None:
        """50% renders as '50%' not '50.0%'."""
        openpyxl = pytest.importorskip("openpyxl")
        wb = openpyxl.Workbook()
        ws = wb.active
        assert ws is not None
        cell = ws.cell(row=1, column=1, value=0.5)
        cell.number_format = "0%"
        wb.save(tmp_path / "whole_pct.xlsx")

        from dd_agents.tools.read_office import read_office

        result = read_office(str(tmp_path / "whole_pct.xlsx"))
        assert result["status"] == "ok"
        assert "50%" in result["content"]
        assert "50.0%" not in result["content"]

    # --- E-2: Unformatted numbers pass through unchanged ---

    def test_plain_numbers_unchanged(self, tmp_path: Path) -> None:
        """Numbers without special format render via str()."""
        openpyxl = pytest.importorskip("openpyxl")
        wb = openpyxl.Workbook()
        ws = wb.active
        assert ws is not None
        ws.cell(row=1, column=1, value=42)
        ws.cell(row=1, column=2, value=3.14)
        wb.save(tmp_path / "plain.xlsx")

        from dd_agents.tools.read_office import read_office

        result = read_office(str(tmp_path / "plain.xlsx"))
        assert result["status"] == "ok"
        assert "42" in result["content"]
        assert "3.14" in result["content"]

    # --- E-3: Sub-table detection ---

    def test_blank_rows_create_sub_tables(self, tmp_path: Path) -> None:
        """Blank rows split the sheet into separate markdown tables."""
        openpyxl = pytest.importorskip("openpyxl")
        wb = openpyxl.Workbook()
        ws = wb.active
        assert ws is not None
        # Sub-table 1: Revenue
        ws.append(["Revenue", "Q1"])
        ws.append(["Product A", 100])
        # Blank separator row
        ws.append([None, None])
        # Sub-table 2: Costs
        ws.append(["Costs", "Q1"])
        ws.append(["Rent", 50])
        wb.save(tmp_path / "subtables.xlsx")

        from dd_agents.tools.read_office import read_office

        result = read_office(str(tmp_path / "subtables.xlsx"))
        assert result["status"] == "ok"
        content = result["content"]

        # Two separate markdown table headers (| A | B | ... | --- | --- |)
        assert content.count("| --- |") == 2

        # All data is preserved
        assert "Revenue" in content
        assert "Product A" in content
        assert "Costs" in content
        assert "Rent" in content

        # Row count excludes blank rows
        assert "4 rows" in content

    def test_no_blank_rows_single_table(self, tmp_path: Path) -> None:
        """Without blank rows, output is a single table (backwards-compatible)."""
        openpyxl = pytest.importorskip("openpyxl")
        wb = openpyxl.Workbook()
        ws = wb.active
        assert ws is not None
        ws.append(["A", 1])
        ws.append(["B", 2])
        ws.append(["C", 3])
        wb.save(tmp_path / "single.xlsx")

        from dd_agents.tools.read_office import read_office

        result = read_office(str(tmp_path / "single.xlsx"))
        assert result["status"] == "ok"
        content = result["content"]

        # Single table header
        assert content.count("| --- |") == 1
        assert "3 rows" in content

    def test_multiple_consecutive_blank_rows_collapsed(self, tmp_path: Path) -> None:
        """Multiple consecutive blank rows count as one boundary."""
        openpyxl = pytest.importorskip("openpyxl")
        wb = openpyxl.Workbook()
        ws = wb.active
        assert ws is not None
        ws.append(["Data1", 10])
        ws.append([None, None])
        ws.append([None, None])
        ws.append([None, None])
        ws.append(["Data2", 20])
        wb.save(tmp_path / "multi_blank.xlsx")

        from dd_agents.tools.read_office import read_office

        result = read_office(str(tmp_path / "multi_blank.xlsx"))
        assert result["status"] == "ok"
        content = result["content"]
        # Still just 2 sub-tables, not 4
        assert content.count("| --- |") == 2
        assert "2 rows" in content


class TestCellEdgeCases:
    """Tests for numeric edge cases in _format_cell."""

    def test_nan_cell_renders_as_string(self) -> None:
        """NaN with currency format falls through to str() instead of crashing."""
        from types import SimpleNamespace

        from dd_agents.tools.read_office import _format_cell

        cell = SimpleNamespace(value=float("nan"), number_format="$#,##0.00")
        result = _format_cell(cell)
        assert "nan" in result.lower()
        # Must NOT produce "$nan" — that would look like a currency value.
        assert "$" not in result

    def test_inf_cell_renders_as_string(self) -> None:
        """Inf with percentage format falls through to str() instead of crashing."""
        from types import SimpleNamespace

        from dd_agents.tools.read_office import _format_cell

        cell = SimpleNamespace(value=float("inf"), number_format="0.0%")
        result = _format_cell(cell)
        assert "inf" in result.lower()
        assert "%" not in result

    def test_negative_currency(self, tmp_path: Path) -> None:
        """Negative currency values render with symbol and negative sign."""
        openpyxl = pytest.importorskip("openpyxl")
        wb = openpyxl.Workbook()
        ws = wb.active
        assert ws is not None
        cell = ws.cell(row=1, column=1, value=-1234.56)
        cell.number_format = "$#,##0.00"
        wb.save(tmp_path / "neg.xlsx")

        from dd_agents.tools.read_office import read_office

        result = read_office(str(tmp_path / "neg.xlsx"))
        assert result["status"] == "ok"
        content = result["content"]
        assert "$" in content
        assert "1,234.56" in content


class TestPathSecurity:
    """Tests for allowed_dir path containment (security boundary)."""

    def test_path_inside_allowed_dir(self, tmp_path: Path) -> None:
        """Files inside allowed_dir are readable."""
        openpyxl = pytest.importorskip("openpyxl")
        wb = openpyxl.Workbook()
        ws = wb.active
        assert ws is not None
        ws.append(["OK", 1])
        wb.save(tmp_path / "inside.xlsx")

        from dd_agents.tools.read_office import read_office

        result = read_office(str(tmp_path / "inside.xlsx"), allowed_dir=str(tmp_path))
        assert result["status"] == "ok"
        assert "OK" in result["content"]

    def test_path_outside_allowed_dir_blocked(self, tmp_path: Path) -> None:
        """Files outside allowed_dir are blocked."""
        openpyxl = pytest.importorskip("openpyxl")
        # Create file in a sibling directory.
        outside = tmp_path / "outside"
        outside.mkdir()
        wb = openpyxl.Workbook()
        ws = wb.active
        assert ws is not None
        ws.append(["Secret", 999])
        wb.save(outside / "secret.xlsx")

        allowed = tmp_path / "allowed"
        allowed.mkdir()

        from dd_agents.tools.read_office import read_office

        result = read_office(str(outside / "secret.xlsx"), allowed_dir=str(allowed))
        assert result["status"] == "error"
        assert "traversal" in result["reason"].lower() or "outside" in result["reason"].lower()

    def test_no_allowed_dir_permits_any_path(self, tmp_path: Path) -> None:
        """Without allowed_dir, any accessible path is permitted."""
        openpyxl = pytest.importorskip("openpyxl")
        wb = openpyxl.Workbook()
        ws = wb.active
        assert ws is not None
        ws.append(["Free", 1])
        wb.save(tmp_path / "free.xlsx")

        from dd_agents.tools.read_office import read_office

        result = read_office(str(tmp_path / "free.xlsx"))
        assert result["status"] == "ok"


class TestReadDocx:
    """Tests for reading Word documents (.docx)."""

    def test_read_docx(self, tmp_path: Path) -> None:
        """Read a .docx file returns text content."""
        pytest.importorskip("markitdown")

        # Create a minimal .docx using python-docx if available, else skip
        try:
            from docx import Document

            doc = Document()
            doc.add_paragraph("This Agreement is between Party A and Party B.")
            doc.add_paragraph("Section 1. Definitions.")
            doc.save(tmp_path / "contract.docx")
        except ImportError:
            pytest.skip("python-docx not available for test fixture creation")

        from dd_agents.tools.read_office import read_office

        result = read_office(str(tmp_path / "contract.docx"))
        assert result["status"] == "ok"
        assert "Party A" in result["content"]
        assert "Definitions" in result["content"]


class TestReadErrors:
    """Tests for error handling."""

    def test_file_not_found(self) -> None:
        """Missing file returns error."""
        from dd_agents.tools.read_office import read_office

        result = read_office("/nonexistent/file.xlsx")
        assert result["status"] == "error"
        assert "not found" in result["reason"].lower() or "not exist" in result["reason"].lower()

    def test_unsupported_extension(self, tmp_path: Path) -> None:
        """Non-Office extension returns error."""
        txt = tmp_path / "notes.txt"
        txt.write_text("hello")

        from dd_agents.tools.read_office import read_office

        result = read_office(str(txt))
        assert result["status"] == "error"
        assert "unsupported" in result["reason"].lower()

    def test_unsupported_pdf(self, tmp_path: Path) -> None:
        """PDF is not handled by read_office — use Read tool instead."""
        pdf = tmp_path / "doc.pdf"
        pdf.write_bytes(b"%PDF-1.4 fake")

        from dd_agents.tools.read_office import read_office

        result = read_office(str(pdf))
        assert result["status"] == "error"

    def test_corrupted_xlsx(self, tmp_path: Path) -> None:
        """Corrupted Excel file returns error, not crash."""
        bad = tmp_path / "corrupt.xlsx"
        bad.write_bytes(b"NOT_AN_EXCEL_FILE")

        from dd_agents.tools.read_office import read_office

        result = read_office(str(bad))
        assert result["status"] == "error"
        assert "reason" in result


class TestOutputTruncation:
    """Tests for output size management."""

    def test_large_xlsx_truncated(self, tmp_path: Path) -> None:
        """Large Excel output is truncated at the limit."""
        openpyxl = pytest.importorskip("openpyxl")
        wb = openpyxl.Workbook()
        ws = wb.active
        assert ws is not None
        ws.append(["Col_A", "Col_B", "Col_C"])
        # Write 5000 rows to generate large output
        for i in range(5000):
            ws.append([f"value_{i}_aaaa", f"data_{i}_bbbb", f"text_{i}_cccc"])
        wb.save(tmp_path / "big.xlsx")

        from dd_agents.tools.read_office import read_office

        result = read_office(str(tmp_path / "big.xlsx"))
        assert result["status"] == "ok"
        assert len(result["content"]) <= 160_000  # Allow small overhead
        assert "truncated" in result["content"].lower()


class TestFallbackToExtractedText:
    """Tests for fallback to pre-extracted markdown in index/text/."""

    def test_fallback_when_primary_read_fails(self, tmp_path: Path) -> None:
        """When primary read fails, falls back to extracted text."""
        # Create an xlsx that will fail to read
        bad_xlsx = tmp_path / "broken.xlsx"
        bad_xlsx.write_bytes(b"PK\x03\x04broken zip content")

        # Create the extracted text fallback
        text_dir = tmp_path / "index" / "text"
        text_dir.mkdir(parents=True)
        (text_dir / "broken.xlsx.md").write_text("# Extracted Content\nRevenue: $1M")

        from dd_agents.tools.read_office import read_office

        result = read_office(str(bad_xlsx), text_dir=str(text_dir))
        assert result["status"] == "ok"
        assert "Revenue" in result["content"]
        assert result.get("method") == "extracted_text_fallback"

    def test_fallback_with_full_path_convention(self, tmp_path: Path) -> None:
        """Fallback resolves via full-path convention (slashes → __)."""
        # Place the corrupt file in a subdirectory to get a multi-segment path
        sub = tmp_path / "data_room"
        sub.mkdir()
        bad_xlsx = sub / "broken.xlsx"
        bad_xlsx.write_bytes(b"PK\x03\x04broken zip content")

        text_dir = tmp_path / "index" / "text"
        text_dir.mkdir(parents=True)

        # Use the canonical _safe_text_name to generate the correct fallback filename
        from dd_agents.extraction.pipeline import ExtractionPipeline

        safe = ExtractionPipeline._safe_text_name(str(bad_xlsx))
        (text_dir / safe).write_text("# Full Path Fallback\nCosts: $2M")

        from dd_agents.tools.read_office import read_office

        result = read_office(str(bad_xlsx), text_dir=str(text_dir))
        assert result["status"] == "ok"
        assert "Costs" in result["content"]
        assert result.get("method") == "extracted_text_fallback"

    def test_markitdown_failure_tries_fallback(self, tmp_path: Path) -> None:
        """When markitdown raises an error, fallback is still attempted."""
        from unittest.mock import patch

        bad_docx = tmp_path / "report.docx"
        bad_docx.write_bytes(b"PK\x03\x04dummy")

        text_dir = tmp_path / "index" / "text"
        text_dir.mkdir(parents=True)
        (text_dir / "report.docx.md").write_text("# Extracted Content\nFallback data")

        from dd_agents.tools.read_office import read_office

        with patch(
            "dd_agents.tools.read_office._read_with_markitdown",
            side_effect=ValueError("markitdown returned empty content"),
        ):
            result = read_office(str(bad_docx), text_dir=str(text_dir))

        assert result["status"] == "ok"
        assert "Fallback data" in result["content"]
        assert result.get("method") == "extracted_text_fallback"

    def test_no_fallback_when_text_dir_missing(self, tmp_path: Path) -> None:
        """Without text_dir, corrupted file returns error."""
        bad = tmp_path / "corrupt.xlsx"
        bad.write_bytes(b"NOT_EXCEL")

        from dd_agents.tools.read_office import read_office

        result = read_office(str(bad))
        assert result["status"] == "error"


class TestColLetter:
    """Tests for _col_letter helper."""

    def test_single_letters(self) -> None:
        from dd_agents.tools.read_office import _col_letter

        assert _col_letter(0) == "A"
        assert _col_letter(25) == "Z"

    def test_double_letters(self) -> None:
        from dd_agents.tools.read_office import _col_letter

        assert _col_letter(26) == "AA"
        assert _col_letter(27) == "AB"
        assert _col_letter(51) == "AZ"
        assert _col_letter(52) == "BA"


class TestToolRegistration:
    """Tests for tool registration in server.py."""

    def test_read_office_in_tool_definitions(self) -> None:
        """read_office appears in the tool definition registry."""
        # Reset singleton to pick up any changes
        import dd_agents.tools.server as srv
        from dd_agents.tools.server import create_tool_definitions

        srv._ALL_TOOL_DEFINITIONS = None

        tools = create_tool_definitions()
        names = [t["name"] for t in tools]
        assert "read_office" in names

    def test_read_office_tool_schema(self) -> None:
        """read_office tool has correct input schema."""
        import dd_agents.tools.server as srv
        from dd_agents.tools.server import create_tool_definitions

        srv._ALL_TOOL_DEFINITIONS = None

        tools = create_tool_definitions()
        tool = next(t for t in tools if t["name"] == "read_office")
        assert "file_path" in tool["input_schema"]["properties"]
        assert "file_path" in tool["input_schema"]["required"]
        assert "sheet_name" in tool["input_schema"]["properties"]

    def test_read_office_in_specialist_custom_tools(self) -> None:
        """read_office in SPECIALIST_CUSTOM_TOOLS list."""
        from dd_agents.tools.server import SPECIALIST_CUSTOM_TOOLS

        assert "read_office" in SPECIALIST_CUSTOM_TOOLS

    def test_read_office_in_specialist_tools(self) -> None:
        """read_office in SPECIALIST_TOOLS list."""
        from dd_agents.agents.specialists import SPECIALIST_TOOLS

        assert "read_office" in SPECIALIST_TOOLS


class TestPromptIntegration:
    """Tests for prompt builder integration."""

    def test_file_access_mentions_read_office(self) -> None:
        """File access instructions mention the read_office tool."""
        from dd_agents.agents.prompt_builder import PromptBuilder

        result = PromptBuilder._build_file_access_instructions()
        assert "read_office" in result

    def test_file_access_mentions_office_extensions(self) -> None:
        """Instructions mention the specific Office extensions."""
        from dd_agents.agents.prompt_builder import PromptBuilder

        result = PromptBuilder._build_file_access_instructions()
        assert ".xlsx" in result
        assert ".docx" in result
